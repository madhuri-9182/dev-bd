import google.generativeai as genai
from pdfminer.high_level import extract_text
from django.conf import settings
from docx import Document
from datetime import datetime
import json
import os
import re


# Load environment variables
genai.configure(api_key=settings.GOOGLE_API_KEY)

# Supported file extensions
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def process_resume(file_path):
    """
    Processes a single resume file and returns parsed data.
    Accepts a file path as input and handles all steps:
    - Text extraction
    - Resume parsing using Gemini API
    - Data validation and formatting
    """

    def allowed_file(filename):
        """Check if the file has an allowed extension."""
        return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS

    def extract_resume_text(file_path):
        """Extracts text from a resume file (PDF, DOCX, or DOC)."""
        try:
            if file_path.lower().endswith(".pdf"):
                return extract_text(file_path)
            elif file_path.lower().endswith(".docx"):
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            elif file_path.lower().endswith(".doc"):
                temp_docx_path = os.path.splitext(file_path)[0] + ".docx"
                convert_doc_to_docx(file_path, temp_docx_path)
                doc = Document(temp_docx_path)
                return "\n".join([para.text for para in doc.paragraphs])
            else:
                print(f"Unsupported file format: {file_path}")
                return ""
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""

    def convert_doc_to_docx(doc_path, docx_path):
        """Converts a .doc file to .docx using unoconv."""
        try:
            import subprocess

            subprocess.run(
                ["unoconv", "-f", "docx", "-o", docx_path, doc_path], check=True
            )
        except Exception as e:
            print(f"Error converting DOC to DOCX: {e}")

    def parse_resume(text):
        """
        Uses Gemini to parse a single resume.
        Returns parsed resume data.
        """
        prompt = (
            "You are an expert resume parser. Extract the following details:\n"
            "1. Name (full name exactly as shown)\n"
            "2. Email (complete address without spaces)\n"
            "3. Phone Number (with country code if available)\n"
            "4. Graduation Date (month and year of last degree completion in 'Month YYYY' format)\n"
            "5. Current Company Name (official legal name)\n"
            "6. Current Designation (exact job title)\n\n"
            "Return STRICT JSON object. The object MUST follow:\n"
            "{\n"
            '    "name": "John Doe",\n'
            '    "email": "john@email.com",\n'
            '    "phoneNumber": "+11234567890",\n'
            '    "graduationDate": "May 2023",\n'
            '    "currentCompanyName": "Tech Corp",\n'
            '    "currentDesignation": "Software Engineer"\n'
            "}\n"
            "Important Rules:\n"
            "- Phone numbers must start with '+' followed by country code\n"
            "- Remove all spaces from emails\n"
            "- Use full month names (January, February etc.)\n"
            "- If information is missing, use 'Not Found'\n"
            "- Current company is the most recent/last mentioned job\n"
        )
        prompt += f"Resume to parse:\n{text}"
        model = genai.GenerativeModel("gemini-pro")
        try:
            response = model.generate_content(prompt)
            raw_response = response.text.strip()
        except Exception as e:
            print(f"Error generating Gemini response: {e}")
            return {}
        # Clean response
        json_start = raw_response.find("{")
        json_end = raw_response.rfind("}") + 1
        if json_start != -1 and json_end != 0:
            raw_response = raw_response[json_start:json_end]
        # Validate JSON
        try:
            parsed_data = json.loads(raw_response)
            return parsed_data
        except Exception as e:
            print(f"JSON parsing failed: {e}")
            print(f"Raw response: {raw_response}")
            return {}

    def calculate_years_of_experience(graduation_date):
        """Calculates experience from graduation date to current date."""
        if not graduation_date or graduation_date == "Not Found":
            return {"year": 0, "month": 0}
        try:
            # Clean up the date string (remove commas and extra spaces)
            cleaned_date = re.sub(r"[^\w\s]", "", graduation_date).strip()
            grad_date = datetime.strptime(cleaned_date, "%B %Y")
            current_date = datetime.now()
            if grad_date > current_date:
                return {"year": 0, "month": 0}
            total_months = (current_date.year - grad_date.year) * 12 + (
                current_date.month - grad_date.month
            )
            if total_months < 0:
                total_months += 12
            return {"year": total_months // 12, "month": total_months % 12}
        except Exception as e:
            print(f"Experience calculation error: {e}")
            return {"year": 0, "month": 0}

    def validate_phone_number(number):
        """Standardizes phone number format."""
        if not number or number == "Not Found":
            return "Not Found"
        cleaned = re.sub(r"(?!^\+)\D", "", number)
        if cleaned.startswith("+"):
            return cleaned
        return f"+{cleaned}" if cleaned else "Not Found"

    # Main processing logic
    file_name = os.path.basename(file_path)
    if not allowed_file(file_name):
        print(f"Unsupported file format: {file_name}")
        return None
    if not os.path.isfile(file_path):
        print(f"File does not exist: {file_path}")
        return None

    # Extract text
    text = extract_resume_text(file_path)
    if not text:
        print(f"Skipped {file_name} (no text extracted)")
        return None

    # Parse resume
    parsed_data = parse_resume(text)

    # Process parsed data
    grad_date = parsed_data.get("graduationDate", "Not Found")
    years_of_experience = calculate_years_of_experience(grad_date)

    processed_data = {
        "file_name": file_name,
        "name": parsed_data.get("name", "Not Found"),
        "email": parsed_data.get("email", "Not Found").replace(" ", ""),
        "phone_number": validate_phone_number(parsed_data.get("phoneNumber")),
        "years_of_experience": years_of_experience,
        "current_company": parsed_data.get("currentCompanyName", "Not Found"),
        "current_designation": parsed_data.get("currentDesignation", "Not Found"),
    }

    return processed_data


# if __name__ == "__main__":
#     # Create a temporary directory for uploaded files
#     with tempfile.TemporaryDirectory() as temp_dir:
#         print(f"Temporary directory created: {temp_dir}")

#         # Prompt user to enter file paths
#         file_paths_input = input("Enter file paths (comma-separated): ").strip()
#         file_paths = [path.strip() for path in file_paths_input.split(",") if path.strip()]

#         if not file_paths:
#             print("No files provided.")
#         else:
#             # Copy files to the temporary directory
#             temp_file_paths = []
#             for file_path in file_paths:
#                 if os.path.isfile(file_path):
#                     temp_file_path = os.path.join(temp_dir, os.path.basename(file_path))
#                     with open(file_path, "rb") as src, open(temp_file_path, "wb") as dst:
#                         dst.write(src.read())
#                     temp_file_paths.append(temp_file_path)
#                 else:
#                     print(f"File does not exist: {file_path}")

#             if temp_file_paths:
#                 # Process each file and collect results
#                 all_parsed_data = []
#                 for temp_file_path in temp_file_paths:
#                     parsed_data = process_resume(temp_file_path)
#                     if parsed_data:
#                         all_parsed_data.append(parsed_data)

#                 # Print results in JSON format
#                 print("\nParsed Resumes:")
#                 print(json.dumps(all_parsed_data, indent=4))
#             else:
#                 print("No valid files to process.")
